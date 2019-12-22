from typing import Tuple, Union

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run



def _insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_paragraph_oxml = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_oxml)
    new_paragraph = Paragraph(new_paragraph_oxml, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style is not None:
        new_paragraph.style = style
    return new_paragraph


class CellObject(object):
    """ An object containning a cell and it's inner:
     - paragraph (holds: runs (w:p element))
     - run (holds: text, pictures, text-styling (font))
     """

    def __init__(self, cell, add_run=True, grid_position=None, paper_size=None,
                 orientation=None):
        self.cell = cell

        cell_paragraph, cell_run = self._get_cell_wrappers(add_run=add_run)
        self.paragraph = cell_paragraph
        self.run = cell_run
        self.grid_position = grid_position
        self.paper_size = paper_size
        self.orientation = orientation

    def _get_cell_wrappers(self, add_run=True) -> Tuple[
        Paragraph, Union[Run, None]]:
        """
        Return the cell's paragraph and create a run object too, return them
        both. They are used to inject elements into the table cell.
        Run object:
        - https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects
        Paragraph Object:
        - https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects
        """
        paragraphs = self.cell.paragraphs
        paragraph = paragraphs[0]
        run = None

        # If we don't add a run then there is no text inserted at the moment
        if add_run:
            run = paragraph.add_run()
        return paragraph, run

    def add_run(self) -> None:
        self.run = self.paragraph.add_run()

    def get_last_paragraph(self) -> Paragraph:
        return self.cell.paragraphs[-1]

    def update_paragraph(self):
        self.paragraph = self.get_last_paragraph()

    def add_paragraph(self, style=None, add_run=True) -> Paragraph:
        self.paragraph = _insert_paragraph_after(self.paragraph, style=style)
        if add_run:
            self.run = self.paragraph.add_run()
        return self.paragraph
