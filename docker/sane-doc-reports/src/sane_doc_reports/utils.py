import base64
import re
import subprocess
import os
import tempfile
from io import BytesIO
import importlib
from pathlib import Path
import traceback
from typing import List
from shutil import which

import arrow
from docx.oxml import OxmlElement
import matplotlib
from docx.text.paragraph import Paragraph
from matplotlib import pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.ticker as mticker

from sane_doc_reports.domain import CellObject, Section
from sane_doc_reports.domain.Section import Section as SectionFactory
from sane_doc_reports.conf import SIZE_H_INCHES, SIZE_W_INCHES, \
    DEFAULT_DPI, DEFAULT_LEGEND_FONT_SIZE, DEFAULT_WORD_FONT, \
    DEFAULT_ALPHA, DEFAULT_FONT_COLOR, DEFAULT_WORD_FONT_FALLBACK, \
    DEFAULT_FONT_AXIS_COLOR, LEGEND_STYLE, DEBUG, WIDTH_POSITION_KEY, \
    HEIGHT_POSITION_KEY, A3_MM_WIDTH, A4_MM_WIDTH, A4_MM_HEIGHT, \
    LETTER_MM_HEIGHT, LETTER_MM_WIDTH, A3_MM_HEIGHT, MAX_AXIS_LABELS, \
    RESIZE_PLOT_ITEMS_AMOUNT_THRESHOLD


def open_b64_image(image_base64):
    """
    Open a virtual image file from base64 format of image.
    """
    prefix_regex = r'^data:.*?;base64,'
    raw_base64 = re.sub(prefix_regex, '', image_base64)
    f = BytesIO()
    f.write(base64.b64decode(raw_base64))
    f.seek(0)
    return f


def fix_svg_to_png(contents):
    if which('svgexport') is None:
        raise Exception('svgexport is not found!')

    tmp_image = open_b64_image(contents)
    tmp_path = '/tmp/_tmp.svg'
    with open(tmp_path, 'wb') as out:
        out.write(tmp_image.read())

    out_path = '/tmp/_out.png'
    out = subprocess.run(['svgexport', tmp_path, out_path], stdout=subprocess.DEVNULL,
                         stderr=subprocess.STDOUT, check=True)

    if DEBUG:
        print("[Sane-doc-reports] Svg conversion output: ", out)

    outf = BytesIO()
    with open(out_path, 'rb') as of:
        outf.write(of.read())

    os.remove(tmp_path)
    os.remove(out_path)
    return outf


def insert_by_type(type: str, cell_object: CellObject,
                   section: Section, trace=False):
    """ Call a elements elemnt's insert method """
    try:
        func = importlib.import_module(f'sane_doc_reports.elements.{type}')
        func.invoke(cell_object, section)
    except ModuleNotFoundError:
        import sane_doc_reports.elements.unimplemented as unimplemented
        unimplemented.invoke(cell_object, section)
    except Exception as e:
        # We want to have a graceful failure instead of early quitting.
        # Maybe we can "salvage" other elements that were generated
        # without any exceptions. Here we will display the faulty
        # elements in the doc.
        if DEBUG:
            traceback.print_exc()
        trace_str = f'\n({traceback.format_exc()})' if trace else ''
        error_msg = f'{section.type} had an error: `{repr(e)}`{trace_str}'
        insert_error(cell_object, error_msg)


def insert_error(cell_object, error_msg):
    from sane_doc_reports.elements import error
    """ Insert an error element """
    section = SectionFactory("error", error_msg, {}, {}, {})
    error.invoke(cell_object, section)


def _insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    return new_para


def add_run(cell_object):
    """ Insert a paragraph so we could add a new element"""
    cell_object.paragraph = _insert_paragraph_after(cell_object.paragraph)
    cell_object.run = cell_object.paragraph.add_run()
    return cell_object


def has_run(cell_object: CellObject):
    """ A helper used to make sure to add a run """
    if cell_object.run is None:
        cell_object.add_run()


def plot(func):
    """ A decorator used to clear and resize each chart """

    def wrapper(*args, **kwargs):
        if plt:
            plt.close()
        plt.clf()
        plt.cla()
        # Fix cropping of plot
        plt.rcParams['figure.constrained_layout.use'] = True
        func(*args, **kwargs)

    return wrapper


def plt_t0_b64(plt: matplotlib.pyplot, figsize=None, dpi=None):
    """ Matplotlib to base64 url """
    path = Path(tempfile.mkdtemp()) / Path(
        next(tempfile._get_candidate_names()) + '.png')

    figsize = figsize if figsize else (1, 1)
    dpi = dpi if dpi else DEFAULT_DPI

    # Remove paddings
    plt.tight_layout()

    plt.savefig(str(path), format='png', figsize=figsize,
                dpi=dpi)

    with open(str(path), "rb") as f:
        img_base64 = base64.b64encode(f.read()).decode("utf-8", "ignore")
        b64 = f'data:image/png;base64,{img_base64}'

    path.unlink()
    return b64


def has_anomalies(items: List):
    return max(items) / min(items) > RESIZE_PLOT_ITEMS_AMOUNT_THRESHOLD


def convert_plt_size(section: Section, cell_object: CellObject,
                     has_anomalies=False):
    """ Convert the plot size from pixels to word """
    size_w, size_h, dpi = (SIZE_W_INCHES, SIZE_H_INCHES, DEFAULT_DPI)

    if WIDTH_POSITION_KEY in section.layout and HEIGHT_POSITION_KEY in section.layout:
        # We need to get the size in inches.
        # w & h are the width in grid size in the word - we need to convert them
        # to inches.
        # ratio in inches: (width_size_in_inches / 12)
        #   Width in inches: w * ratio_in_inches

        sizes = {
            'A4': {
                'portrait': A4_MM_WIDTH.inches,
                'landscape': A4_MM_HEIGHT.inches
            },
            'A3': {
                'portrait': A3_MM_WIDTH.inches,
                'landscape': A3_MM_HEIGHT.inches
            },
            'LETTER': {
                'portrait': LETTER_MM_WIDTH.inches,
                'landscape': LETTER_MM_HEIGHT.inches
            }
        }
        page_size = sizes['A4'] if not cell_object.paper_size else \
            sizes[cell_object.paper_size]
        page_orientation = 'portrait' if not cell_object.orientation else \
            cell_object.orientation

        width_size_in_inches = page_size[page_orientation]
        ratio_w = width_size_in_inches / 12
        w = int(section.layout[WIDTH_POSITION_KEY])
        size_w = (ratio_w * w)

    if has_anomalies:
        size_w += 1
        size_h += 1

    return size_w, size_h, dpi


def get_ax_location(legend_style):
    """ Get the legend location from the verticalAlign key or return default """
    align = legend_style.get('align', None)
    vertical_align = legend_style.get('verticalAlign', None)

    if not align or not vertical_align:
        return 'best'

    vertical_align = vertical_align.replace('top', 'upper').replace(
        'bottom', 'lower')
    return f'{vertical_align} {align}'


def get_current_li(extra, list_type='List Number'):
    """ Return the current list item style and indent level """
    list_type = list_type if 'list_type' not in extra else extra['list_type']
    if not extra or 'list_level' not in extra:
        return list_type, 0, list_type

    extra_list_level = int(extra['list_level'])
    list_level = 0
    if extra_list_level == 0:
        list_level = 2
        p_style = list_type
    elif extra_list_level > 3:
        # The docx template doesn't support more than
        #   4 levels of indentation.
        list_level = 4
        p_style = f'{list_type} {list_level}'
    elif extra_list_level > 0:
        list_level += extra['list_level'] + 1
        p_style = f'{list_type} {list_level}'

    return p_style, list_level, list_type


def list_number(doc, par, prev=None, level=None, num=True):
    """
    Taken from: https://github.com/python-openxml/python-docx/issues/25
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def remove_plot_borders(ax):
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)


def set_axis_font(ax):
    font = fm.FontProperties(family=get_chart_font(),
                             size=DEFAULT_LEGEND_FONT_SIZE)

    ax.tick_params(axis='x', colors=DEFAULT_FONT_AXIS_COLOR)
    ax.tick_params(axis='y', colors=DEFAULT_FONT_AXIS_COLOR)

    for label in ax.get_xticklabels():
        label.set_fontproperties(font)

    for label in ax.get_yticklabels():
        label.set_fontproperties(font)


def set_legend_max_count(ax, cell_object: CellObject):
    g_pos = cell_object.grid_position
    width_ratio = g_pos['width'] / g_pos['global_cols']
    axis_count = MAX_AXIS_LABELS
    if width_ratio <= 0.2:
        axis_count = MAX_AXIS_LABELS / 4
    elif width_ratio < 0.6:
        axis_count = MAX_AXIS_LABELS / 2

    myLocator = mticker.MaxNLocator(axis_count)
    ax.xaxis.set_major_locator(myLocator)


def set_legend_style(legend, options=None):
    plt.gcf().autofmt_xdate()
    if options:
        if 'hideLegend' in options and options['hideLegend']:
            plt.gca().legend().set_visible(False)
            return

    legend.get_frame().set_alpha(DEFAULT_ALPHA)
    legend.get_frame().set_linewidth(0.0)

    font = fm.FontProperties(family=get_chart_font(),
                             size=DEFAULT_LEGEND_FONT_SIZE)

    for text in legend.get_texts():
        text.set_fontproperties(font)
        text.set_color(DEFAULT_FONT_COLOR)
        if 'valign' in options:
            text.set_position((0, options['valign']))


def change_legend_vertical_alignment(section: Section, top=0):
    section.layout[LEGEND_STYLE]['valign'] = top
    return section


def get_chart_font():
    names = [f.name for f in matplotlib.font_manager.fontManager.ttflist]

    if DEFAULT_WORD_FONT not in names:
        return DEFAULT_WORD_FONT_FALLBACK
    return DEFAULT_WORD_FONT


def get_formatted_date(input_date,
                       layout=None) -> str:
    """ Returns the formatted date string
            input_date - date we want to format
            layout - custom formats from the sane JSONs

        Note: ParserError is raised and should be catched if used.
    """
    date = arrow.now()

    # Use the date if supplied, and not now()
    if input_date:
        date = arrow.get(input_date)

    formatted_date = date.isoformat()

    # Use the user supplied format
    if layout and 'format' in layout:
        formatted_date = date.format(layout['format'])

    return formatted_date
