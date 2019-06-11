from typing import Union

from docx.table import _Cell

from sane_doc_reports.conf import STYLE_KEY
from sane_doc_reports.domain.CellObject import CellObject
from sane_doc_reports.domain.Section import Section
from sane_doc_reports.elements import text
from sane_doc_reports.styles.utils import apply_style
from sane_doc_reports.utils import has_run


def insert_elem(cell_object: Union[CellObject, _Cell], section: Section,
                add_run=False):
    """ Insert text into a specified cell, can add a style and a run element
        too.
     """

    # Get the relevant cell object
    if isinstance(cell_object, _Cell):
        cell_object = CellObject(cell_object, add_run=add_run)
    elif add_run:
        cell_object = cell_object
        cell_object.add_run()

    if not add_run:
        has_run(cell_object)

    # Apply the relevant style
    apply_style(cell_object, section)

    text.invoke(cell_object, section, apply_styling=False)


def insert_text(cell_object: Union[CellObject, _Cell],
                section: Union[str, Section],
                style={}, add_run=False):

    if isinstance(section, str):
        section = Section('text', section, {}, {})
        section.set_style(style)

    section.type = 'text'
    section.add_style(style)
    insert_elem(cell_object, section, add_run)


def insert_header(cell_object: Union[CellObject, _Cell],
                  section: Union[str, Section],
                  style={}, add_run=False, header='h1'):

    if isinstance(section, str):
        section = Section('header', section, {}, {})
        section.set_style(style)

    section.extra = {**{'header_tag': header}, **section.extra}
    section.add_style(style)
    insert_elem(cell_object, section, add_run)
