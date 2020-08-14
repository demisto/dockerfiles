import os

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Inches

from sane_doc_reports.domain.CellObject import CellObject
from sane_doc_reports.domain.Page import Page
from sane_doc_reports.domain.Section import Section
from sane_doc_reports.domain import SaneJson
from sane_doc_reports.elements import image
from sane_doc_reports.utils import insert_by_type
from sane_doc_reports.conf import DEBUG, A4_MM_HEIGHT, A4_MM_WIDTH, \
    TOP_MARGIN_PT, BOTTOM_MARGIN_PT, LEFT_MARGIN_PT, RIGHT_MARGIN_PT, \
    A3_MM_WIDTH, A3_MM_HEIGHT, LETTER_MM_WIDTH, LETTER_MM_HEIGHT, PAPER_A4, \
    PAPER_A3, PAPER_LETTER, DOCX_TEMAPLTE_FILE, XSOAR_LOGO_BASE64, \
    MAX_CUSTOMER_LOGO_WIDTH_INCH, MAX_CUSTOMER_LOGO_HEIGHT_INCH
from sane_doc_reports.populate.grid import get_cell, merge_cells
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT


def _debug_show_styles(document):
    styles = document.styles
    styles_p = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    styles_t = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
    styles = styles_p + styles_t
    print("Styles: ")
    for style in styles:
        print("\t", style.name)


class Report:
    """
    In charge of generating a DOCX report form a SANE report (JSON)
    """

    def __init__(self, pages: List[Page], sane_json: SaneJson, options={}):
        template_path = Path(os.path.dirname(__file__)) / DOCX_TEMAPLTE_FILE
        with template_path.open('rb') as f:
            self.document = Document(f)

            # Remove the default paragraph in the template.
            self.document._body.clear_content()

        self.pages = pages

        # Used to calculate and create the page grid(layout)
        self.sane_json = sane_json

        self.options = options
        self.page_width = A4_MM_WIDTH
        self.page_height = A4_MM_HEIGHT

    def populate_report(self) -> None:
        if not self.options.get('disableHeaders', False):
            self.add_header_logos()

        paper_size = self.options.get('paper_size', 'A4')
        self.change_page_size(paper_size)
        self._decrease_layout_margins()
        page_count = self.sane_json.get_pages_count() - 1
        orientation = self.options.get('orientation', 'portrait')

        for page_num, sane_page in enumerate(self.sane_json.get_sane_pages()):
            cols, rows = sane_page.calculate_page_grid()

            if DEBUG:
                print(f'Creating a layout grid of size ({rows},{cols})' +
                      f' for page: {page_num}')

            if orientation == 'landscape':
                if DEBUG:
                    print("Changing orientation to landscape.")
                self.orient_landscape()

            grid = self.document.add_table(rows=rows, cols=cols)
            if DEBUG:
                grid.style = 'Table Grid'

            page = self.pages[page_num]
            for section in page:
                cell, grid_pos = get_cell(grid, section)
                grid_pos = {
                    "width": 0 if not len(grid_pos) == 2 else grid_pos[0],
                    "height": 0 if not len(grid_pos) == 2 else grid_pos[1],
                    "global_rows": rows,
                    "global_cols": cols
                }
                merge_cells(grid, section)
                cell_object = CellObject(cell, add_run=False,
                                         grid_position=grid_pos,
                                         paper_size=paper_size,
                                         orientation=orientation)
                self._insert_section(cell_object, section)

            # If this isn't the last page, we can add another page break.
            if page_num != page_count:
                p = self.document.add_paragraph()
                r = p.add_run()
                if DEBUG:
                    r.text = f'Page break ({page_num})'
                r.add_break(WD_BREAK.PAGE)

    def _insert_section(self, cell_object: CellObject,
                        section: Section) -> None:
        section_type = section.type

        # Fix the chart name
        if section_type == 'chart':
            section_type = section.layout['chartType'] + '_chart'
            section.type = section_type

        insert_by_type(section_type, cell_object, section,
                       trace=self.options.get('trace', False))

    def save(self, output_file_path: str):
        self.document.save(output_file_path)

    def orient_landscape(self):
        sections = self.document.sections
        for section in sections:
            section.orientation = WD_ORIENT.LANDSCAPE

            # Swap the width and height to get landscape
            section.page_height = self.page_width
            section.page_width = self.page_height

    def change_page_size(self, paper_size: str) -> None:
        if DEBUG:
            print("Paper size:", paper_size)

        if paper_size == PAPER_A4:
            self.page_width = A4_MM_WIDTH
            self.page_height = A4_MM_HEIGHT

        elif paper_size == PAPER_A3:
            self.page_width = A3_MM_WIDTH
            self.page_height = A3_MM_HEIGHT

        elif paper_size == PAPER_LETTER:
            self.page_width = LETTER_MM_WIDTH
            self.page_height = LETTER_MM_HEIGHT

        sections = self.document.sections
        for section in sections:
            section.page_height = self.page_height
            section.page_width = self.page_width

    def _decrease_layout_margins(self) -> None:
        sections = self.document.sections
        for section in sections:
            section.top_margin = Pt(TOP_MARGIN_PT)
            section.bottom_margin = Pt(BOTTOM_MARGIN_PT)
            section.left_margin = Pt(LEFT_MARGIN_PT)
            section.right_margin = Pt(RIGHT_MARGIN_PT)

    def add_header_logos(self):
        # Find the headers
        section = self.document.sections[0]
        section.header_distance = Pt(0)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Inches(24))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # Add the left cell to the header
        left_image = CellObject(left_cell)
        left_cell.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        left_cell.vertical_alignment = 1

        # Add the right cell to the header
        right_image = CellObject(right_cell)
        right_cell.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        right_cell.vertical_alignment = 1

        # Add the main logo
        left_logo_b64 = self.options.get('demistoLogo', XSOAR_LOGO_BASE64)
        s = Section('image', left_logo_b64, {}, {})
        image.invoke(left_image, s)

        # Add the customer logo
        right_logo_b64 = self.options.get('customerLogo', False)
        if right_logo_b64:
            s = Section('image', right_logo_b64, {}, {
                'max_size': {'height': MAX_CUSTOMER_LOGO_HEIGHT_INCH, # max size in inches
                             'width': MAX_CUSTOMER_LOGO_WIDTH_INCH}})
            image.invoke(right_image, s)
