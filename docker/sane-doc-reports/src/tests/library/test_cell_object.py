from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from sane_doc_reports.domain.CellObject import CellObject


def test_cell_object_init():
    d = Document()
    table = d.add_table(1, 1)
    cell = table.cell(0, 0)
    co = CellObject(cell)
    assert co.cell == cell
    assert isinstance(co.paragraph, Paragraph)
    assert isinstance(co.run, Run)


def test_cell_object_add_run():
    d = Document()
    table = d.add_table(1, 1)
    cell = table.cell(0, 0)
    co = CellObject(cell)
    cur_run = co.run
    co.add_run()
    assert co.run != cur_run


def test_cell_object_add_paragraph():
    d = Document()
    table = d.add_table(1, 1)
    cell = table.cell(0, 0)
    co = CellObject(cell)
    co.run.text = 'old'

    assert len(d.element.xpath("//w:t[text()='old']")) == 1
    assert len(d.element.xpath("//w:t[text()='new']")) == 0

    co.add_paragraph()
    co.run.text = 'new'

    assert len(d.element.xpath("//w:t[text()='old']")) == 1
    assert len(d.element.xpath("//w:t[text()='new']")) == 1

    # Here we check it acctually came after
    check_order = d.element.xpath(
        '//w:p//w:t[text()="new"]/preceding::w:p//w:t[text()="old"]')
    assert check_order[0].text == 'old'


def test_cell_object_get_last_paragraph():
    d = Document()
    table = d.add_table(1, 1)
    cell = table.cell(0, 0)
    co = CellObject(cell)
    co.run.text = 'old'
    p = co.add_paragraph()
    co.run.text = 'new'

    last_p = co.get_last_paragraph()
    last_p.text = 'newnew'
    assert p.text == 'newnew'
    assert co.paragraph.text == 'newnew'

